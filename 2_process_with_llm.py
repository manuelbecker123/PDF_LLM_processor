# 2_process_with_llm.py
import os
import re
import time
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor
import config # Import the configuration

# Initialize OpenAI client
# Try to get API key from environment variable first
api_key_to_use = os.environ.get("OPENAI_API_KEY")
if not api_key_to_use and config.OPENAI_API_KEY:
    print("INFO: Using OpenAI API Key from config.py. "
          "It's recommended to set it as an environment variable (OPENAI_API_KEY).")
    api_key_to_use = config.OPENAI_API_KEY
elif not api_key_to_use and not config.OPENAI_API_KEY:
    print("ERROR: OpenAI API Key not found.")
    print("Please set it as an environment variable OPENAI_API_KEY or in config.py (OPENAI_API_KEY).")
    exit()

client = OpenAI(api_key=api_key_to_use)

# Global variable for the task instruction
TASK_INSTRUCTION = ""

def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def load_prompt_template(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def save_text_as_word(text_content, output_file_path, is_original=False):
    doc = Document()
    page_pattern = r"(={50}\nSTART OF PAGE: \d+\n={50})(.*?)(={50}\nEND OF PAGE: \d+\n={50})"
    pages = re.findall(page_pattern, text_content, re.DOTALL)

    for i, page_parts in enumerate(pages):
        page_header, page_content, page_footer = page_parts

        if i > 0:
            doc.add_page_break()

        # Add header, content, footer
        p_header = doc.add_paragraph(page_header.strip())
        p_content = doc.add_paragraph(page_content.strip())
        p_footer = doc.add_paragraph(page_footer.strip())

        # Font styling (optional, can be customized)
        sentences = re.split(r'[.!?]', page_content.strip())
        num_sentences = len([s for s in sentences if s.strip()])

        font_size = Pt(config.DEFAULT_FONT_SIZE_PT)
        if num_sentences > config.SENTENCE_THRESHOLD_FOR_SMALLER_FONT:
            font_size = Pt(config.SMALLER_FONT_SIZE_PT)

        for p in [p_header, p_content, p_footer]:
            for run in p.runs:
                run.font.name = 'Calibri' # Or your preferred font
                run.font.size = font_size
                if is_original and "(IMAGE:" in run.text: # Highlight image refs in original
                     run.font.highlight_color = WD_COLOR_INDEX.YELLOW


    doc.save(output_file_path)
    status_type = "Original" if is_original else "Processed"
    print(f"{status_type} Word document saved as {output_file_path}")

def call_llm_for_page(page_text_content, page_number_str, base_prompt_template):
    global TASK_INSTRUCTION # Use the global task instruction

    prompt = base_prompt_template.replace("{{PAGE_TEXT}}", page_text_content)
    prompt = prompt.replace("{{TASK_INSTRUCTION}}", TASK_INSTRUCTION)

    retries = 3
    delay = 20  # seconds
    for attempt in range(retries):
        try:
            completion = client.chat.completions.create(
                model=config.OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}]
            )
            return completion.choices[0].message.content
        except Exception as e:
            print(f"Error processing page {page_number_str} with OpenAI: {e}.")
            if "rate limit" in str(e).lower() or "quota" in str(e).lower() or "overloaded" in str(e).lower() :
                print(f"Rate limit or quota issue. Retrying attempt {attempt + 1}/{retries} after {delay} seconds...")
                time.sleep(delay)
                delay *= 2 # Exponential backoff
            else:
                return f"ERROR PROCESSING PAGE {page_number_str}: {str(e)}" # Non-retryable error for this page
    return f"ERROR: Failed to process page {page_number_str} after {retries} retries due to API issues."


def process_chapter_text_file(extracted_text_filepath, chapter_name, output_chapter_dir, base_prompt_template):
    print(f"\nProcessing text file for chapter: {chapter_name}")
    extracted_text = read_text_file(extracted_text_filepath)

    page_pattern = r"(={50}\nSTART OF PAGE: (\d+)\n={50})(.*?)(={50}\nEND OF PAGE: \2\n={50})"
    pages_data = re.findall(page_pattern, extracted_text, re.DOTALL)

    processed_full_text = ""
    total_pages = len(pages_data)

    # Using ThreadPoolExecutor for concurrent page processing
    # Note: OpenAI API calls are I/O bound, so threading can help.
    # Be mindful of API rate limits.
    with ThreadPoolExecutor(max_workers=config.MAX_WORKERS) as executor:
        future_to_page = {}
        for i, page_match in enumerate(pages_data):
            page_header, page_number_str, page_content_raw, page_footer = page_match
            print(f"Submitting page {page_number_str}/{total_pages} from {chapter_name} for processing...")
            # Submit to executor: function, arg1, arg2, ...
            future = executor.submit(call_llm_for_page, page_content_raw.strip(), page_number_str, base_prompt_template)
            future_to_page[future] = (page_header, page_number_str, page_footer, i) # Store original order index

        # Collect results in order
        results = [None] * len(pages_data)
        for future in future_to_page:
            page_header, page_number_str, page_footer, original_index = future_to_page[future]
            try:
                processed_page_content = future.result()
                print(f"Received processed content for page {page_number_str}/{total_pages} from {chapter_name}.")
                # Reconstruct the page with processed content
                results[original_index] = f"{page_header}\n{processed_page_content.strip()}\n{page_footer}\n\n"
            except Exception as exc:
                print(f"Page {page_number_str} generated an exception: {exc}")
                results[original_index] = f"{page_header}\nERROR: Content generation failed.\n{page_footer}\n\n"

        processed_full_text = "".join(results)


    # Save the processed content as a Word document
    processed_docx_path = os.path.join(output_chapter_dir, f"processed_{chapter_name}.docx")
    save_text_as_word(processed_full_text, processed_docx_path, is_original=False)

    # Save the original content as a separate Word document for comparison
    original_docx_path = os.path.join(output_chapter_dir, f"original_{chapter_name}.docx")
    save_text_as_word(extracted_text, original_docx_path, is_original=True)

    # Copy the images folder from the source to the output chapter directory
    source_images_folder = os.path.join(os.path.dirname(extracted_text_filepath), f"{chapter_name}_imgs")
    if os.path.exists(source_images_folder):
        output_images_folder = os.path.join(output_chapter_dir, f"{chapter_name}_imgs")
        if os.path.exists(output_images_folder):
            shutil.rmtree(output_images_folder) # Remove if exists to avoid error on copytree
        shutil.copytree(source_images_folder, output_images_folder)
        print(f"Copied images from {source_images_folder} to {output_images_folder}")
    else:
        print(f"No images folder found at {source_images_folder} for chapter {chapter_name}.")


def main():
    global TASK_INSTRUCTION # To modify the global variable

    if not os.path.exists(config.INPUT_EXTRACTED_FOLDER):
        print(f"ERROR: Input folder with extracted text not found: {config.INPUT_EXTRACTED_FOLDER}")
        print("Please run '1_extract_content.py' first or update 'INPUT_EXTRACTED_FOLDER' in config.py.")
        return

    if not os.path.exists(config.PROCESSED_OUTPUT_FOLDER):
        os.makedirs(config.PROCESSED_OUTPUT_FOLDER)
        print(f"Created output folder: {config.PROCESSED_OUTPUT_FOLDER}")

    try:
        base_prompt_template = load_prompt_template(config.PROMPT_FILE_PATH)
    except FileNotFoundError:
        print(f"ERROR: Prompt template file not found at {config.PROMPT_FILE_PATH}")
        print("Please create this file or update its path in config.py.")
        return

    print("\n--- LLM Page Processor ---")
    print("The script will process text from each page using an LLM based on your prompt.")
    print(f"Using OpenAI Model: {config.OPENAI_MODEL}")
    print(f"Prompt template loaded from: {config.PROMPT_FILE_PATH}")
    
    # Get the specific task instruction from the user
    print("\nPlease specify the core task for the LLM (e.g., 'Translate to German', 'Summarize in 3 bullet points', 'Generate 3 quiz questions and answers'):")
    TASK_INSTRUCTION = input("Enter your task instruction: ").strip()
    if not TASK_INSTRUCTION:
        print("No task instruction provided. Exiting.")
        return
    print(f"Task instruction set to: '{TASK_INSTRUCTION}'")


    chapter_folders = [d for d in os.listdir(config.INPUT_EXTRACTED_FOLDER)
                       if os.path.isdir(os.path.join(config.INPUT_EXTRACTED_FOLDER, d))
                       and not d.endswith("_imgs")] # Avoid processing image folders directly

    if not chapter_folders:
        print(f"No chapter folders found in {config.INPUT_EXTRACTED_FOLDER}. Ensure '1_extract_content.py' has been run.")
        return

    print(f"Found {len(chapter_folders)} chapter(s) to process.")

    for chapter_name in chapter_folders:
        extracted_text_file = os.path.join(config.INPUT_EXTRACTED_FOLDER, chapter_name, f"extracted_text_{chapter_name}.txt")
        output_chapter_directory = os.path.join(config.PROCESSED_OUTPUT_FOLDER, chapter_name)

        if not os.path.exists(extracted_text_file):
            print(f"Skipping {chapter_name}: Text file not found at {extracted_text_file}")
            continue

        os.makedirs(output_chapter_directory, exist_ok=True)
        process_chapter_text_file(extracted_text_file, chapter_name, output_chapter_directory, base_prompt_template)

    print("\nAll LLM processing finished.")
    print(f"Processed output can be found in: {config.PROCESSED_OUTPUT_FOLDER}")

if __name__ == '__main__':
    main()